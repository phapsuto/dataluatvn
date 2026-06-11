import os
# Configure OpenMP and thread settings to prevent macOS crashes
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["OPENBLAS_NUM_THREADS"] = "1"
os.environ["VECLIB_MAXIMUM_THREADS"] = "1"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["HF_HUB_OFFLINE"] = "1"
os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["DISABLE_LLM_EXPANSION"] = "1"

import sqlite3
import io
import re
from datetime import datetime
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from fastapi.responses import StreamingResponse

from fastapi import APIRouter, Depends, Query, Path, HTTPException

from app.dependencies import require_api_key
from app.database import get_db_connection, get_content_connection, simple_ttl_cache
from app.schemas.laws import (
    StatsResponse, PaginatedSearchResponse, CategoryItem,
    LawDetail, RelationshipInfo, ArticleModification,
    ProvinceItem, WardItem, PaginatedChunkSearchResponse,
)

router = APIRouter(prefix="/laws", tags=["🔍 Tìm kiếm & Tra cứu (Luật)"])


# ─────────────────── STATISTICS ───────────────────

@simple_ttl_cache(ttl_seconds=3600)
def _get_cached_stats():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT count(*) FROM documents")
    total_docs = cursor.fetchone()[0]

    cursor.execute("SELECT count(*) FROM relationships")
    total_rels = cursor.fetchone()[0]

    cursor.execute("SELECT loai_van_ban, count(*) FROM documents GROUP BY loai_van_ban ORDER BY count(*) DESC")
    types_count = {row[0] or "Khác": row[1] for row in cursor.fetchall()}

    cursor.execute("SELECT tinh_trang_hieu_luc, count(*) FROM documents GROUP BY tinh_trang_hieu_luc ORDER BY count(*) DESC")
    status_count = {row[0] or "Không xác định": row[1] for row in cursor.fetchall()}

    conn.close()
    return {
        "total_documents": total_docs,
        "total_relationships": total_rels,
        "by_document_type": types_count,
        "by_effectiveness_status": status_count,
    }


@router.get("/stats", response_model=StatsResponse, tags=["📊 Thống kê (Luật)"], summary="Thống kê tổng quan")
def get_stats(_key=Depends(require_api_key)):
    """Thống kê tổng quan cơ sở dữ liệu. **Yêu cầu API Key.**"""
    return _get_cached_stats()


# ─────────────────── SEARCH & RETRIEVAL ───────────────────

# Stopwords tiếng Việt — loại bỏ khi xây dựng FTS query
VIETNAMESE_STOPWORDS = {
    # Đại từ & chỉ định
    "tôi", "tui", "mình", "ta", "chúng", "các", "những", "một", "này", "đó", "kia",
    "nào", "gì", "nào", "ai", "đâu", "sao", "bao", "mấy",
    # Giới từ & liên từ
    "của", "và", "với", "trong", "trên", "dưới", "ngoài", "giữa", "bên", "về",
    "cho", "đến", "từ", "tại", "theo", "bằng", "qua", "vào", "hay", "hoặc",
    "nhưng", "mà", "rằng", "nếu", "khi", "vì", "do", "bởi", "để",
    # Phó từ & trạng từ
    "đã", "đang", "sẽ", "sắp", "vẫn", "còn", "rất", "lắm", "quá", "hơn",
    "nhất", "không", "chưa", "chẳng", "nên", "cần", "phải", "được", "bị",
    "có", "là", "thì", "cũng", "lại", "rồi", "mới", "cứ", "đều",
    # Từ nghi vấn phổ biến
    "thế", "như", "thế nào", "như thế nào", "bao nhiêu", "vậy",
    # Từ chức năng khác
    "việc", "cái", "con", "người", "điều", "khoản",
}

def generate_spelling_variants(word: str) -> List[str]:
    variants = {word}
    replacements = [
        ('òa', 'o\u00e0'), ('óa', 'o\u00e1'), ('ỏa', 'o\u1ea3'), ('õa', 'o\u00e3'), ('ọa', 'o\u1ea1'),
        ('òe', 'o\u00e8'), ('óe', 'o\u00e9'), ('ỏe', 'o\u1ebd'), ('õe', 'o\u1ebd'), ('ọe', 'o\u1eb9'),
        ('ủy', 'u\u1ef7'), ('úy', 'u\u00fd'), ('ùy', 'u\u1ef3'), ('ũy', 'u\u1ef5'), ('ụy', 'u\u1ef9'),
        ('o\u00e0', 'òa'), ('o\u00e1', 'óa'), ('o\u1ea3', 'ỏa'), ('o\u00e3', 'õa'), ('o\u1ea1', 'ọa'),
        ('o\u00e8', 'òe'), ('o\u00e9', 'óe'), ('o\u1ebd', 'ỏe'), ('o\u1ebd', 'õe'), ('o\u1eb9', 'ọe'),
        ('u\u1ef7', 'ủy'), ('u\u00fd', 'úy'), ('u\u1ef3', 'ùy'), ('u\u1ef5', 'ũy'), ('u\u1ef9', 'ụy'),
    ]
    word_lower = word.lower()
    for new_style, old_style in replacements:
        if new_style in word_lower:
            variants.add(word_lower.replace(new_style, old_style))
    return list(variants)

def parse_fts_query(q: str) -> str:
    """
    Phân tích query FTS5 thông minh cho cả keyword search lẫn câu hỏi tự nhiên,
    kết hợp tự động sinh các biến thể dấu tiếng Việt cũ/mới để tối đa hóa Recall.
    """
    # Bước 1: Loại bỏ ký tự đặc biệt FTS
    q_clean = re.sub(r'[^\w\s]', ' ', q).strip()
    all_words = [w for w in q_clean.split() if w]
    if not all_words:
        return ""
    
    # Bước 2: Loại bỏ stopwords và từ quá ngắn (≤1 ký tự)
    keywords = [w for w in all_words if w.lower() not in VIETNAMESE_STOPWORDS and len(w) > 1]
    
    # Nếu sau khi lọc không còn từ nào, dùng lại all_words
    if not keywords:
        keywords = [w for w in all_words if len(w) > 1]
    if not keywords:
        keywords = all_words
        
    def format_keyword(w: str, append_star: bool = False) -> str:
        vars = generate_spelling_variants(w)
        star = "*" if append_star else ""
        if len(vars) > 1:
            return "(" + " OR ".join([f"{v}{star}" for v in vars]) + ")"
        return f"{w}{star}"
    
    # Bước 3: Áp dụng chiến lược theo độ dài
    if len(keywords) == 1:
        return format_keyword(keywords[0], append_star=True)
    
    if len(keywords) <= 3:
        # Query ngắn → AND prefix giữa các từ
        and_query = " AND ".join([format_keyword(w, append_star=True) for w in keywords])
        return and_query
    
    if len(keywords) <= 6:
        # Query vừa → AND tất cả
        and_query = " AND ".join([format_keyword(w, append_star=False) for w in keywords])
        
        # Fallback: AND chỉ top 3 từ dài nhất
        top_words = sorted(keywords, key=len, reverse=True)[:3]
        and_fallback = " AND ".join([format_keyword(w, append_star=False) for w in top_words])
        return f"({and_query}) OR ({and_fallback})"
    
    # Query dài (câu hỏi tự nhiên)
    top_keywords = sorted(keywords, key=len, reverse=True)[:7]
    top_formatted_7 = []
    for w in top_keywords:
        vars = generate_spelling_variants(w)
        if len(vars) > 1:
            top_formatted_7.append("(" + " OR ".join(vars) + ")")
        else:
            top_formatted_7.append(w)
            
    and_top5 = " AND ".join(top_formatted_7[:5])
    and_top3 = " AND ".join(top_formatted_7[:3])
    return f"({and_top5}) OR ({and_top3})"

def get_province_search_terms(province_code: str, conn) -> list:
    cursor = conn.cursor()
    cursor.execute("SELECT name, full_name FROM provinces WHERE code = ?", (province_code,))
    row = cursor.fetchone()
    if not row:
        return []
    p_name = row['name']
    p_fullname = row['full_name']
    terms = {p_name, p_fullname}
    for term in list(terms):
        term_alt = term.replace("oà", "òa").replace("oá", "óa").replace("oả", "ỏa")
        term_alt2 = term.replace("òa", "oà").replace("óa", "oá").replace("ỏa", "oả")
        terms.add(term_alt)
        terms.add(term_alt2)
    return list(terms)

def get_ward_search_terms(ward_code: str, conn) -> list:
    cursor = conn.cursor()
    cursor.execute("SELECT name, full_name FROM wards WHERE code = ?", (ward_code,))
    row = cursor.fetchone()
    if not row:
        return []
    w_name = row['name']
    w_fullname = row['full_name']
    terms = {w_name, w_fullname}
    for term in list(terms):
        term_alt = term.replace("oà", "òa").replace("oá", "óa").replace("oả", "ỏa")
        term_alt2 = term.replace("òa", "oà").replace("óa", "oá").replace("ỏa", "oả")
        terms.add(term_alt)
        terms.add(term_alt2)
    return list(terms)


def extract_type_and_title(q: str):
    q_clean = q.strip().lower()
    types = {
        'hiến pháp': 'Hiến pháp',
        'bộ luật': 'Bộ luật',
        'luật': 'Luật',
        'pháp lệnh': 'Pháp lệnh',
        'nghị định': 'Nghị định',
        'nghị quyết': 'Nghị quyết',
        'quyết định': 'Quyết định',
        'thông tư': 'Thông tư',
        'chỉ thị': 'Chỉ thị',
        'lệnh': 'Lệnh'
    }
    for t_key, t_val in types.items():
        if q_clean.startswith(t_key + ' '):
            return t_val, q_clean[len(t_key) + 1:].strip()
    return None, None


@router.get("/search", response_model=PaginatedSearchResponse, summary="Tìm kiếm văn bản")
def search_laws(
    q: Optional[str] = Query(None, description="Từ khóa tìm kiếm"),
    loai_van_ban: Optional[str] = Query(None, description="Lọc theo loại văn bản"),
    co_quan_ban_hanh: Optional[str] = Query(None, description="Lọc theo cơ quan ban hành"),
    status: Optional[str] = Query(None, alias="tinh_trang", description="Lọc theo tình trạng hiệu lực"),
    linh_vuc: Optional[str] = Query(None, description="Lọc theo lĩnh vực"),
    province_code: Optional[str] = Query(None, description="Mã tỉnh/thành phố để lọc địa phương"),
    ward_code: Optional[str] = Query(None, description="Mã quận huyện/phường xã để lọc địa phương"),
    limit: int = Query(20, ge=1, le=100, description="Số lượng tối đa (1–100)"),
    offset: int = Query(0, ge=0, description="Vị trí bắt đầu"),
    require_content: bool = Query(False, description="Chỉ trả về văn bản có nội dung HTML"),
    _key=Depends(require_api_key),
):
    """
    Tìm kiếm và lọc văn bản pháp luật bằng **Full-Text Search (FTS5)**.

    **Tính năng nổi bật:**
    - Kết quả tự động sắp xếp theo độ liên quan (relevance) nếu có từ khóa `q`.
    - Phân trang đầy đủ với `total_pages`, `current_page`, `has_next`, `has_previous`.

    **Yêu cầu API Key.**
    """
    conn = get_db_connection()
    cursor = conn.cursor()

    where_clauses = []
    params: list = []
    from_clause = "documents d"

    # Filter out future dates (e.g., typos like 3024 or 2055) relative to today's date
    today_str = datetime.now().strftime("%Y-%m-%d")
    where_clauses.append(
        "(d.ngay_ban_hanh IS NULL OR d.ngay_ban_hanh = '' OR "
        "(length(d.ngay_ban_hanh) = 10 AND "
        "substr(d.ngay_ban_hanh, 7, 4) || '-' || substr(d.ngay_ban_hanh, 4, 2) || '-' || substr(d.ngay_ban_hanh, 1, 2) <= ?))"
    )
    params.append(today_str)

    if require_content:
        where_clauses.append("d.has_content = 1")

    subquery_params = []
    q_has_fts = False
    if q:
        fts_query = parse_fts_query(q)
        if fts_query:
            q_has_fts = True
            try:
                cursor.execute("SELECT 1 FROM content_fts LIMIT 1")
                fts_table = "content_fts"
            except Exception:
                fts_table = "documents_fts"
            
            t_val, title_part = extract_type_and_title(q)
            if t_val and title_part:
                subquery_sql = f"""
                  SELECT rowid as id, rank as f_rank FROM {fts_table} WHERE {fts_table} MATCH ?
                  UNION ALL
                  SELECT id, 0 as f_rank FROM documents WHERE so_ky_hieu = ? OR (loai_van_ban = ? AND LOWER(title) LIKE ?)
                """
                subquery_params = [fts_query, q.strip(), t_val, f"%{title_part}%"]
            else:
                subquery_sql = f"""
                  SELECT rowid as id, rank as f_rank FROM {fts_table} WHERE {fts_table} MATCH ?
                  UNION ALL
                  SELECT id, 0 as f_rank FROM documents WHERE so_ky_hieu = ?
                """
                subquery_params = [fts_query, q.strip()]
            
            from_clause = f"({subquery_sql}) u JOIN documents d ON u.id = d.id"

    if province_code:
        terms = get_province_search_terms(province_code, conn)
        if terms:
            province_clauses = []
            for t in terms:
                province_clauses.append("d.pham_vi LIKE ? OR d.co_quan_ban_hanh LIKE ?")
                params.extend([f"%{t}%", f"%{t}%"])
            where_clauses.append(f"({' OR '.join(province_clauses)})")

    if ward_code:
        terms = get_ward_search_terms(ward_code, conn)
        if terms:
            ward_clauses = []
            for t in terms:
                ward_clauses.append("d.pham_vi LIKE ? OR d.co_quan_ban_hanh LIKE ?")
                params.extend([f"%{t}%", f"%{t}%"])
            where_clauses.append(f"({' OR '.join(ward_clauses)})")

    if loai_van_ban:
        where_clauses.append("d.loai_van_ban = ?")
        params.append(loai_van_ban)
    if co_quan_ban_hanh:
        where_clauses.append("d.co_quan_ban_hanh LIKE ?")
        params.append(f"%{co_quan_ban_hanh}%")
    if status:
        where_clauses.append("d.tinh_trang_hieu_luc = ?")
        params.append(status)
    if linh_vuc:
        where_clauses.append("(d.linh_vuc LIKE ? OR d.nganh LIKE ?)")
        params.extend([f"%{linh_vuc}%", f"%{linh_vuc}%"])

    if not where_clauses:
        where_clauses.append("1=1")

    where_sql = " AND ".join(where_clauses)

    if q_has_fts:
        cursor.execute(f"SELECT count(distinct d.id) FROM {from_clause} WHERE {where_sql}", subquery_params + params)
    else:
        cursor.execute(f"SELECT count(*) FROM {from_clause} WHERE {where_sql}", params)
    total = cursor.fetchone()[0]

    order_params = []
    if q_has_fts:
        # Hướng 1: Ưu tiên Thứ bậc Hiệu lực + Trạng thái hoạt động, sau đó mới dùng FTS rank
        order_clause = (
            "CASE WHEN d.so_ky_hieu = ? THEN 2 WHEN d.so_ky_hieu LIKE ? THEN 1 ELSE 0 END DESC, "
            "CASE "
            "  WHEN replace(replace(LOWER(d.title), char(13), ''), char(10), ' ') = ? THEN 100 "
            "  WHEN replace(replace(LOWER(d.loai_van_ban || ' ' || d.title), char(13), ''), char(10), ' ') = ? THEN 100 "
            "  WHEN replace(replace(LOWER(d.loai_van_ban || ' ' || d.title), char(13), ''), char(10), ' ') LIKE ? THEN 50 "
            "  ELSE 0 "
            "END DESC, "
            "CASE "
            "  WHEN d.tinh_trang_hieu_luc LIKE '%Còn hiệu lực%' THEN 3 "
            "  WHEN d.tinh_trang_hieu_luc LIKE '%Hết hiệu lực một phần%' THEN 2 "
            "  WHEN d.tinh_trang_hieu_luc IS NULL OR d.tinh_trang_hieu_luc = '' THEN 1 "
            "  ELSE 0 "
            "END DESC, "
            "CASE LOWER(d.loai_van_ban) "
            "  WHEN 'hiến pháp' THEN 10 "
            "  WHEN 'bộ luật' THEN 9 "
            "  WHEN 'luật' THEN 9 "
            "  WHEN 'pháp lệnh' THEN 8 "
            "  WHEN 'nghị định' THEN 7 "
            "  WHEN 'nghị quyết' THEN 6 "
            "  WHEN 'quyết định' THEN 5 "
            "  WHEN 'thông tư' THEN 4 "
            "  ELSE 1 "
            "END DESC, "
            "CAST(substr(d.ngay_ban_hanh, 7, 4) AS INTEGER) DESC, "
            "u.f_rank, "
            "substr(d.ngay_ban_hanh, 4, 2) DESC, substr(d.ngay_ban_hanh, 1, 2) DESC, d.id DESC"
        )
        q_strip = q.strip()
        q_lower = q_strip.lower()
        order_params = [
            q_strip, f"%{q_strip}%", 
            q_lower, q_lower, f"%{q_lower}%"
        ]
    else:
        order_clause = (
            "CASE WHEN d.tinh_trang_hieu_luc IN ('Hết hiệu lực toàn bộ', 'Hết hiệu lực') THEN 0 ELSE 1 END DESC, "
            "CASE LOWER(d.loai_van_ban) "
            "  WHEN 'hiến pháp' THEN 10 "
            "  WHEN 'bộ luật' THEN 9 "
            "  WHEN 'luật' THEN 9 "
            "  WHEN 'pháp lệnh' THEN 8 "
            "  WHEN 'nghị định' THEN 7 "
            "  WHEN 'nghị quyết' THEN 6 "
            "  WHEN 'quyết định' THEN 5 "
            "  WHEN 'thông tư' THEN 4 "
            "  ELSE 1 "
            "END DESC, "
            "substr(d.ngay_ban_hanh, 7, 4) DESC, substr(d.ngay_ban_hanh, 4, 2) DESC, substr(d.ngay_ban_hanh, 1, 2) DESC, d.id DESC"
        )

    if q_has_fts:
        query_sql = (
            f"SELECT d.id, d.title, d.so_ky_hieu, d.ngay_ban_hanh, d.loai_van_ban, d.co_quan_ban_hanh, d.tinh_trang_hieu_luc, MIN(u.f_rank) as f_rank "
            f"FROM {from_clause} WHERE {where_sql} GROUP BY d.id ORDER BY {order_clause} LIMIT ? OFFSET ?"
        )
        query_params = subquery_params + params + order_params + [limit, offset]
    else:
        query_sql = (
            f"SELECT d.id, d.title, d.so_ky_hieu, d.ngay_ban_hanh, d.loai_van_ban, d.co_quan_ban_hanh, d.tinh_trang_hieu_luc "
            f"FROM {from_clause} WHERE {where_sql} ORDER BY {order_clause} LIMIT ? OFFSET ?"
        )
        query_params = params + [limit, offset]

    cursor.execute(query_sql, query_params)
    rows = cursor.fetchall()
    conn.close()

    total_pages = (total + limit - 1) // limit if limit > 0 else 0
    current_page = (offset // limit) + 1 if limit > 0 else 1
    has_next = current_page < total_pages
    has_previous = current_page > 1

    return {
        "total": total,
        "limit": limit,
        "offset": offset,
        "total_pages": total_pages,
        "current_page": current_page,
        "has_next": has_next,
        "has_previous": has_previous,
        "results": [dict(r) for r in rows],
    }



@router.get("/chunk-search", response_model=PaginatedChunkSearchResponse, summary="Tìm kiếm theo điều khoản (FTS5 trên Chunks)")
def chunk_search_laws(
    q: str = Query(..., description="Từ khóa tìm kiếm"),
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    _key=Depends(require_api_key),
):
    """
    Tìm kiếm chi tiết đến cấp độ Điều (Article) của văn bản sử dụng FTS5 trên bảng chunks.
    Trả về danh sách các chunks (điều khoản) khớp kèm theo thông tin văn bản gốc.
    """
    conn = get_db_connection()
    cursor = conn.cursor()

    fts_query = parse_fts_query(q)
    if not fts_query:
        return {
            "total": 0,
            "limit": limit,
            "offset": offset,
            "total_pages": 0,
            "current_page": 1,
            "has_next": False,
            "has_previous": False,
            "results": [],
        }

    # Count total matching chunks
    try:
        cursor.execute("""
            SELECT count(*) 
            FROM chunks_fts 
            WHERE chunks_fts MATCH ?
        """, (fts_query,))
        total = cursor.fetchone()[0]
    except Exception:
        total = 0

    if total == 0:
        conn.close()
        return {
            "total": 0,
            "limit": limit,
            "offset": offset,
            "total_pages": 0,
            "current_page": 1,
            "has_next": False,
            "has_previous": False,
            "results": [],
        }

    # Fetch matching chunks with metadata
    query = """
        SELECT c.id, c.doc_id, c.chunk_index, c.chunk_type, c.chunk_header, c.chunk_text, c.token_estimate,
               d.title as document_title, d.so_ky_hieu as document_so_ky_hieu, d.loai_van_ban as document_loai_van_ban,
               d.co_quan_ban_hanh as document_co_quan_ban_hanh, d.tinh_trang_hieu_luc as document_tinh_trang_hieu_luc,
               d.ngay_ban_hanh as document_ngay_ban_hanh
        FROM chunks_fts f
        JOIN document_chunks c ON f.rowid = c.id
        JOIN documents d ON c.doc_id = d.id
        WHERE f.chunks_fts MATCH ?
        ORDER BY f.rank
        LIMIT ? OFFSET ?
    """
    
    try:
        cursor.execute(query, (fts_query, limit, offset))
        rows = cursor.fetchall()
    except Exception as e:
        rows = []
        
    conn.close()

    total_pages = (total + limit - 1) // limit if limit > 0 else 0
    current_page = (offset // limit) + 1 if limit > 0 else 1
    has_next = current_page < total_pages
    has_previous = current_page > 1

    return {
        "total": total,
        "limit": limit,
        "offset": offset,
        "total_pages": total_pages,
        "current_page": current_page,
        "has_next": has_next,
        "has_previous": has_previous,
        "results": [dict(r) for r in rows],
    }


# ─────────────────── HYBRID SEARCH ───────────────────

@router.get("/hybrid-search", tags=["🔍 Tìm kiếm & Tra cứu (Luật)"], summary="Tìm kiếm Hybrid (BM25 + FTS5 + RRF)")
async def hybrid_search(
    q: str = Query(..., description="Từ khoá tìm kiếm"),
    limit: int = Query(10, ge=1, le=50),
    _api_key: str = Depends(require_api_key),
):
    """
    Tìm kiếm Hybrid: BM25Okapi + FTS5 + Reciprocal Rank Fusion.
    
    Cho kết quả chính xác hơn nhiều so với FTS5 thuần, đặc biệt với
    câu hỏi ngôn ngữ tự nhiên về pháp luật.
    """
    from app.hybrid_search import get_hybrid_engine
    
    engine = get_hybrid_engine()
    if not engine or not engine.is_ready:
        # Fallback về FTS5 nếu hybrid chưa sẵn sàng
        return await search_laws(q=q, limit=limit, _api_key=_api_key)
    
    results = engine.search(q, top_k=limit)
    
    return {
        "total": len(results),
        "limit": limit,
        "offset": 0,
        "engine": "hybrid_bm25_fts5_rrf",
        "results": results,
    }


# ─────────────────── SMART HYBRID SEARCH (PHASE 3) ───────────────────

def normalize_spelling(text: str) -> str:
    if not text:
        return ""
    text = text.replace('òa', 'o\u00e0').replace('óa', 'o\u00e1').replace('ỏa', 'o\u1ea3').replace('õa', 'o\u00e3').replace('ọa', 'o\u1ea1')
    text = text.replace('òe', 'o\u00e8').replace('óe', 'o\u00e9').replace('ỏe', 'o\u1ebd').replace('õe', 'o\u1ebd').replace('ọe', 'o\u1eb9')
    text = text.replace('ủy', 'u\u1ef7').replace('úy', 'u\u00fd').replace('ùy', 'u\u1ef3').replace('ũy', 'u\u1ef5').replace('ụy', 'u\u1ef9')
    text = text.replace('Òa', 'O\u00e0').replace('Óa', 'O\u00e1').replace('Ỏa', 'O\u1ea3').replace('Õa', 'O\u00e3').replace('Ọa', 'O\u1ea1')
    text = text.replace('Òe', 'O\u00e8').replace('Óe', 'O\u00e9').replace('Ỏe', 'O\u1ebd').replace('Õe', 'O\u1ebd').replace('Ọe', 'O\u1eb9')
    text = text.replace('Ủy', 'U\u1ef7').replace('Úy', 'U\u00fd').replace('Ùy', 'U\u1ef3').replace('Ũy', 'U\u1ef5').replace('Ụy', 'U\u1ef9')
    return text

_SMART_SEARCH_MODEL = None
_SMART_SEARCH_INDEX = None
_SMART_SEARCH_ID_MAP = None
_SMART_SEARCH_PROVINCES = None

def get_smart_search_provinces(conn):
    global _SMART_SEARCH_PROVINCES
    if _SMART_SEARCH_PROVINCES is None:
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT name, full_name FROM provinces")
            rows = cursor.fetchall()
            prov_set = set()
            for name, full_name in rows:
                prov_set.add(name.lower())
                prov_set.add(full_name.lower())
                # Tên viết tắt hoặc bỏ chữ "Tỉnh", "Thành phố"
                short = name.replace("Tỉnh", "").replace("Thành phố", "").strip().lower()
                if short:
                    prov_set.add(short)
            _SMART_SEARCH_PROVINCES = prov_set
        except Exception as e:
            print(f"⚠️ Không thể load provinces: {e}")
            _SMART_SEARCH_PROVINCES = set()
    return _SMART_SEARCH_PROVINCES

def get_smart_search_resources():
    global _SMART_SEARCH_MODEL, _SMART_SEARCH_INDEX, _SMART_SEARCH_ID_MAP
    import os
    from app.config import EMBEDDING_MODEL_SOTA, FAISS_INDEX_SOTA
    
    # Lazy load sentence-transformers & PyTorch
    if _SMART_SEARCH_MODEL is None:
        try:
            import torch
            from sentence_transformers import SentenceTransformer
            device = "mps" if torch.backends.mps.is_available() else ("cuda" if torch.cuda.is_available() else "cpu")
            model_kwargs = {"torch_dtype": torch.float16} if device in ["mps", "cuda"] else {}
            _SMART_SEARCH_MODEL = SentenceTransformer(EMBEDDING_MODEL_SOTA, device=device, model_kwargs=model_kwargs)
            print(f"✅ Loaded Embedding Model {EMBEDDING_MODEL_SOTA} on {device.upper()} (float16)")
        except Exception as e:
            print(f"⚠️ Không thể load embedding model {EMBEDDING_MODEL_SOTA}: {e}")
            
    # Lazy load FAISS index
    if _SMART_SEARCH_INDEX is None:
        index_path = FAISS_INDEX_SOTA
        if os.path.exists(index_path):
            try:
                import faiss
                import numpy as np
                _SMART_SEARCH_INDEX = faiss.read_index(index_path)
                if hasattr(_SMART_SEARCH_INDEX, "id_map"):
                    id_arr = faiss.vector_to_array(_SMART_SEARCH_INDEX.id_map)
                    _SMART_SEARCH_ID_MAP = {int(cid): i for i, cid in enumerate(id_arr)}
                else:
                    _SMART_SEARCH_ID_MAP = {}
                
                # Cấu hình nprobe cho các loại chỉ mục IVF (Inverted File)
                if hasattr(_SMART_SEARCH_INDEX, "nprobe"):
                    nprobe_val = int(os.environ.get("FAISS_NPROBE", "64"))
                    _SMART_SEARCH_INDEX.nprobe = nprobe_val
                    print(f"🎯 Chỉ mục FAISS loại IVF được thiết lập nprobe = {nprobe_val}")
            except Exception as e:
                print(f"⚠️ Không thể load FAISS index từ {index_path}: {e}")
        else:
            print(f"⚠️ Không tìm thấy file index: {index_path}")
            
    return _SMART_SEARCH_MODEL, _SMART_SEARCH_INDEX

_RERANKER_MODEL = None
_RERANKER_TOKENIZER = None

def get_vietnamese_reranker():
    global _RERANKER_MODEL, _RERANKER_TOKENIZER
    import os
    from app.config import RERANKER_MODEL_SOTA
    
    if _RERANKER_MODEL is None:
        try:
            import torch
            from transformers import AutoModelForSequenceClassification, AutoTokenizer
            device = "mps" if torch.backends.mps.is_available() else ("cuda" if torch.cuda.is_available() else "cpu")
            _RERANKER_TOKENIZER = AutoTokenizer.from_pretrained(RERANKER_MODEL_SOTA)
            # Load in float16 on MPS/CUDA for low latency
            if device in ["mps", "cuda"]:
                _RERANKER_MODEL = AutoModelForSequenceClassification.from_pretrained(
                    RERANKER_MODEL_SOTA, torch_dtype=torch.float16
                ).to(device)
            else:
                _RERANKER_MODEL = AutoModelForSequenceClassification.from_pretrained(RERANKER_MODEL_SOTA).to(device)
            _RERANKER_MODEL.eval()
            print(f"✅ Loaded Vietnamese Reranker: {RERANKER_MODEL_SOTA} on device: {device.upper()} (float16)")
        except Exception as e:
            print(f"⚠️ Không thể load Cross-Encoder Reranker {RERANKER_MODEL_SOTA}: {e}")
    return _RERANKER_MODEL, _RERANKER_TOKENIZER


def preprocess_and_correct_query(q: str) -> str:
    """
    Tiền xử lý câu truy vấn: Khôi phục dấu tiếng Việt (nếu viết không dấu)
    và sửa lỗi chính tả bằng cách gọi LLM FPT Cloud (có cache & timeout).
    """
    import sqlite3
    import requests
    import re
    from app.config import MEMORY_DB
    
    clean_q = q.strip()
    if not clean_q:
        return ""
        
    # Lọc bỏ nếu là số ký hiệu thuần túy
    if re.search(r'\b\d+/', clean_q):
        return clean_q
        
    # Tra cache trước
    try:
        conn = sqlite3.connect(MEMORY_DB)
        cursor = conn.cursor()
        cursor.execute("CREATE TABLE IF NOT EXISTS query_preprocess_cache (query_text TEXT PRIMARY KEY, processed_text TEXT NOT NULL)")
        conn.commit()
        
        cursor.execute("SELECT processed_text FROM query_preprocess_cache WHERE query_text = ?", (clean_q.lower(),))
        row = cursor.fetchone()
        conn.close()
        if row:
            return row[0]
    except Exception as e:
        print(f"⚠️ Lỗi đọc cache preprocess: {e}")

    # Chỉ gọi LLM khi câu dài hơn 2 từ
    words = clean_q.split()
    if len(words) <= 2:
        return clean_q

    from app.config import FPT_CLOUD_API_KEY
    if not FPT_CLOUD_API_KEY or os.environ.get("DISABLE_LLM_EXPANSION") == "1":
        return clean_q

    FPT_URL = "https://mkp-api.fptcloud.com/v1/chat/completions"
    FPT_MODEL = "gemma-4-31B-it"
    headers = {
        "Authorization": f"Bearer {FPT_CLOUD_API_KEY}",
        "Content-Type": "application/json"
    }

    system_prompt = (
        "Bạn là chuyên gia ngôn ngữ tiếng Việt pháp luật.\n"
        "Nhiệm vụ của bạn là đọc câu hỏi của người dùng, khôi phục dấu tiếng Việt đầy đủ nếu câu hỏi viết không dấu, và sửa lỗi chính tả nếu có.\n"
        "Quy tắc tuyệt đối:\n"
        "- Trả về duy nhất câu đã được khôi phục dấu và sửa lỗi.\n"
        "- Giữ nguyên các thuật ngữ pháp lý.\n"
        "- KHÔNG viết thêm bất kỳ lời giải thích, mở đầu hoặc kết luận nào."
    )

    payload = {
        "model": FPT_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": "thu tuc ly hon va chia dat"},
            {"role": "assistant", "content": "thủ tục ly hôn và chia đất"},
            {"role": "user", "content": "mức phạt vi phạp giao thông"},
            {"role": "assistant", "content": "mức phạt vi phạm giao thông"},
            {"role": "user", "content": clean_q}
        ],
        "temperature": 0.1,
        "max_tokens": 100
    }

    try:
        response = requests.post(FPT_URL, json=payload, headers=headers, timeout=1.5)
        if response.status_code == 200:
            data = response.json()
            processed = data["choices"][0]["message"]["content"].strip()
            processed = processed.strip('"').strip("'")
            if processed:
                # Ghi vào cache
                try:
                    conn = sqlite3.connect(MEMORY_DB)
                    cursor = conn.cursor()
                    cursor.execute("INSERT OR REPLACE INTO query_preprocess_cache (query_text, processed_text) VALUES (?, ?)", (clean_q.lower(), processed))
                    conn.commit()
                    conn.close()
                except Exception:
                    pass
                return processed
    except Exception as e:
        print(f"⚠️ Lỗi gọi API preprocess: {e}")

    return clean_q


@router.get("/smart-search", response_model=PaginatedChunkSearchResponse, summary="Tìm kiếm Hybrid thông minh (FTS5 Chunks + FAISS Vector Chunks + RRF + Metadata Boost)")
def smart_search_laws(
    q: str = Query(..., description="Từ khóa tìm kiếm tự nhiên"),
    loai_van_ban: Optional[str] = Query(None, description="Lọc theo loại văn bản"),
    co_quan_ban_hanh: Optional[str] = Query(None, description="Lọc theo cơ quan ban hành"),
    status: Optional[str] = Query(None, alias="tinh_trang", description="Lọc theo tình trạng hiệu lực"),
    linh_vuc: Optional[str] = Query(None, description="Lọc theo lĩnh vực"),
    limit: int = Query(10, ge=1, le=50),
    offset: int = Query(0, ge=0),
    _key=Depends(require_api_key),
):
    """
    Tìm kiếm ngữ nghĩa kết hợp FTS5, FAISS Vector và thuật toán RRF Boosting (khớp số ký hiệu, độ hiệu lực pháp lý, tình trạng văn bản).
    **Yêu cầu API Key.**
    """
    # Phòng thủ nếu hàm được gọi trực tiếp bằng Python (không qua FastAPI)
    from fastapi.params import Query as FastAPIQuery, Depends as FastAPIDepends
    
    def is_fastapi_param(val):
        t_str = str(type(val))
        return "Query" in t_str or "Depends" in t_str or isinstance(val, (FastAPIQuery, FastAPIDepends))
        
    if is_fastapi_param(loai_van_ban): loai_van_ban = None
    if is_fastapi_param(co_quan_ban_hanh): co_quan_ban_hanh = None
    if is_fastapi_param(status): status = None
    if is_fastapi_param(linh_vuc): linh_vuc = None
    if is_fastapi_param(limit): limit = 10
    if is_fastapi_param(offset): offset = 0

    if not q.strip():
        return {
            "total": 0, "limit": limit, "offset": offset,
            "total_pages": 0, "current_page": 1,
            "has_next": False, "has_previous": False, "results": []
        }

    # Tiền xử lý khôi phục dấu và sửa lỗi chính tả bằng LLM
    q_processed = preprocess_and_correct_query(q)

    conn = get_db_connection()
    cursor = conn.cursor()

    # Load danh sách tỉnh thành phục vụ nhận diện địa phương
    PROVINCES = get_smart_search_provinces(conn)
    q_lower = q_processed.lower()
    has_locality_in_q = False
    for p in PROVINCES:
        if p in q_lower:
            has_locality_in_q = True
            break

    # Xây dựng các điều kiện lọc (SQL Filter Clauses)
    where_clauses = []
    sql_params = []
    
    if loai_van_ban:
        where_clauses.append("d.loai_van_ban = ?")
        sql_params.append(loai_van_ban)
    if co_quan_ban_hanh:
        where_clauses.append("d.co_quan_ban_hanh LIKE ?")
        sql_params.append(f"%{co_quan_ban_hanh}%")
    if status:
        where_clauses.append("d.tinh_trang_hieu_luc = ?")
        sql_params.append(status)
    if linh_vuc:
        where_clauses.append("(d.linh_vuc LIKE ? OR d.nganh LIKE ?)")
        sql_params.extend([f"%{linh_vuc}%", f"%{linh_vuc}%"])

    # 1. Trích xuất số ký hiệu linh hoạt và phân tích ý định câu hỏi
    # A. Trích xuất số ký hiệu đầy đủ có gạch chéo (ví dụ: 15/2020/NĐ-CP)
    symbol_patterns = re.findall(r'\b\d+/(?:[A-Za-z0-9À-ỹ-]+/)*[A-Za-z0-9À-ỹ-]+\b', q_processed)
    extracted_symbols = [s.strip().lower() for s in symbol_patterns]

    # B. Trích xuất các số hiệu đơn lẻ (ví dụ: số 151, số 180, số 15)
    numbers_in_q = []
    for num in re.findall(r'\b\d+\b', q_processed):
        val = int(num)
        if not (2000 <= val <= 2030):
            numbers_in_q.append(num)

    # C. Phân tích ý định câu hỏi (Dynamic Intent Detection)
    is_detail_query = False
    detail_keywords = [
        "thủ tục", "trình tự", "hồ sơ", "mức phạt", "xử phạt", "phạt bao nhiêu",
        "thời hiệu", "thẩm quyền", "điều kiện để", "cấp phép", "giấy phép",
        "báo cáo", "hạn nộp", "nộp ở đâu", "cơ quan nào có thẩm quyền", "ai có thẩm quyền"
    ]
    for kw in detail_keywords:
        if kw in q_lower:
            is_detail_query = True
            break
            
    is_general_query = False
    general_keywords = [
        "khái niệm", "định nghĩa", "nguyên tắc", "chính sách", "quy định chung",
        "bộ luật", "luật đất đai", "luật hình sự", "luật dân sự", "luật bảo hiểm"
    ]
    for kw in general_keywords:
        if kw in q_lower:
            is_general_query = True
            break

    # Gọi Query Expansion
    from app.query_expansion import expand_query
    expanded_terms = expand_query(q_processed)

    # 2. Chạy FTS5 Search trên Chunks (Top 100)
    q_clean = re.sub(r'[^\w\s]', ' ', q_processed).strip()
    words_in_q = [w for w in q_clean.split() if w]
    keywords_in_q = [w for w in words_in_q if w.lower() not in VIETNAMESE_STOPWORDS and len(w) > 1]
    
    run_fts = False
    if len(extracted_symbols) > 0 or len(keywords_in_q) <= 3:
        run_fts = True

    fts_results = []
    if run_fts:
        fts_query = parse_fts_query(q_processed)
        if fts_query:
            if expanded_terms:
                expanded_clauses = [f'"{term}"' for term in expanded_terms if term.strip()]
                if expanded_clauses:
                    fts_query = f"({fts_query}) OR " + " OR ".join(expanded_clauses)
                    
            fts_symbols = []
            if extracted_symbols:
                fts_symbols.extend([re.sub(r'[^\w\s]', ' ', sym).strip() for sym in extracted_symbols])
            if numbers_in_q:
                fts_symbols.extend(numbers_in_q)
                
            if fts_symbols:
                symbol_or_clause = " OR ".join([f'"{s}"' for s in fts_symbols if s])
                if symbol_or_clause:
                    fts_query = f"({symbol_or_clause}) OR ({fts_query})"

        if fts_query:
            try:
                fts_where = " AND ".join(["f.chunks_fts MATCH ?"] + where_clauses)
                fts_sql_params = [fts_query] + sql_params
                cursor.execute(f"""
                    SELECT c.id, c.doc_id, f.rank as fts_rank
                    FROM chunks_fts f
                    CROSS JOIN document_chunks c ON f.rowid = c.id
                    JOIN documents d ON c.doc_id = d.id
                    WHERE {fts_where}
                    ORDER BY f.rank
                    LIMIT 100
                """, fts_sql_params)
                fts_results = [dict(row) for row in cursor.fetchall()]
            except Exception as e:
                print(f"⚠️ Lỗi truy vấn FTS5 Chunks: {e}")

    query_vector = None
    # 3. Chạy Vector Search trên FAISS (Top 150)
    vector_results = []
    model, faiss_index = get_smart_search_resources()
    
    if model and faiss_index:
        try:
            import numpy as np
            import faiss
            q_norm = normalize_spelling(q_processed)
            if expanded_terms:
                q_norm = q_norm + " " + " ".join([normalize_spelling(term) for term in expanded_terms])
                
            query_vector = model.encode([q_norm], show_progress_bar=False, convert_to_numpy=True)
            query_vector = query_vector.astype(np.float32)
            faiss.normalize_L2(query_vector)
            
            distances, indices = faiss_index.search(query_vector, 150)
            
            for score, cid in zip(distances[0], indices[0]):
                if cid != -1:
                    vector_results.append((int(cid), float(score)))
        except Exception as e:
            print(f"⚠️ Lỗi truy vấn Vector Search: {e}")

    # Lọc Post-filtering trong SQL cho kết quả FAISS
    vector_matched = []
    cid_to_doc = {}
    cid_to_vec_score = {}
    if vector_results:
        vector_ids = [item[0] for item in vector_results]
        cid_to_vec_score = {item[0]: item[1] for item in vector_results}
        placeholders = ",".join(["?"] * len(vector_ids))
        vec_where = " AND ".join([f"c.id IN ({placeholders})"] + where_clauses)
        vec_sql_params = list(vector_ids) + sql_params
        try:
            cursor.execute(f"""
                SELECT c.id, c.doc_id 
                FROM document_chunks c
                JOIN documents d ON c.doc_id = d.id
                WHERE {vec_where}
            """, vec_sql_params)
            for row in cursor.fetchall():
                cid_to_doc[row["id"]] = row["doc_id"]
            
            vector_matched = [cid for cid in vector_ids if cid in cid_to_doc]
        except Exception as e:
            print(f"⚠️ Lỗi lọc FAISS results: {e}")

    # 4. Normalized Score Fusion
    # A. FTS5 min-max normalization
    fts_scores = {item["id"]: item["fts_rank"] for item in fts_results}
    if fts_scores:
        fts_vals = list(fts_scores.values())
        min_fts = min(fts_vals)
        max_fts = max(fts_vals)
        fts_range = max_fts - min_fts
        if fts_range < 1e-6:
            fts_range = 1.0
        fts_norm = {cid: (max_fts - val) / fts_range for cid, val in fts_scores.items()}
    else:
        fts_norm = {}

    # B. FAISS min-max normalization
    if vector_matched:
        vec_vals = [cid_to_vec_score[cid] for cid in vector_matched]
        min_vec = min(vec_vals)
        max_vec = max(vec_vals)
        vec_range = max_vec - min_vec
        if vec_range < 1e-6:
            vec_range = 1.0
        vec_norm = {cid: (cid_to_vec_score[cid] - min_vec) / vec_range for cid in vector_matched}
    else:
        vec_norm = {}

    # C. Linear Fusion (0.3 * Sparse + 0.7 * Dense)
    fused_scores = {}
    
    for item in fts_results:
        cid = item["id"]
        doc_id = item["doc_id"]
        fused_scores[cid] = {
            "id": cid,
            "doc_id": doc_id,
            "score": 0.3 * fts_norm[cid]
        }
        
    for cid in vector_matched:
        doc_id = cid_to_doc[cid]
        if cid not in fused_scores:
            fused_scores[cid] = {
                "id": cid,
                "doc_id": doc_id,
                "score": 0.0
            }
        fused_scores[cid]["score"] += 0.7 * vec_norm[cid]

    # Sắp xếp các ứng viên theo điểm fused score
    sorted_candidates = sorted(fused_scores.values(), key=lambda x: x["score"], reverse=True)
    
    # Lấy Top ứng viên để chạy Reranker (Cấu hình linh hoạt cho Latency/Recall)
    rerank_limit = int(os.environ.get("RERANK_LIMIT", "15"))
    candidates_to_rerank = sorted_candidates[:rerank_limit]
    remaining_candidates = sorted_candidates[rerank_limit:]

    # D. Chạy Vietnamese Cross-Encoder Reranker
    if candidates_to_rerank and os.environ.get("DISABLE_RERANKER") != "1":
        candidate_ids = [item["id"] for item in candidates_to_rerank]
        placeholders = ",".join(["?"] * len(candidate_ids))
        
        try:
            cursor.execute(f"SELECT id, chunk_text FROM document_chunks WHERE id IN ({placeholders})", candidate_ids)
            chunk_texts_map = {row[0]: row[1] for row in cursor.fetchall()}
            
            rerank_model, rerank_tokenizer = get_vietnamese_reranker()
            if rerank_model is not None and rerank_tokenizer is not None:
                import torch
                pairs = [[q_processed, chunk_texts_map.get(cid, "")] for cid in candidate_ids]
                
                rerank_max_length = int(os.environ.get("RERANK_MAX_LENGTH", "256"))
                inputs = rerank_tokenizer(
                    pairs, 
                    padding=True, 
                    truncation=True, 
                    max_length=rerank_max_length, 
                    return_tensors="pt"
                )
                inputs = {k: v.to(rerank_model.device) for k, v in inputs.items()}
                
                with torch.no_grad():
                    outputs = rerank_model(**inputs)
                    logits = outputs.logits.cpu().numpy()
                    
                    if len(logits.shape) > 1 and logits.shape[1] > 1:
                        raw_scores = logits[:, 1].tolist()
                    else:
                        raw_scores = logits.squeeze(-1).tolist()
                        if not isinstance(raw_scores, list):
                            raw_scores = [raw_scores]
                            
                min_raw = min(raw_scores)
                max_raw = max(raw_scores)
                raw_range = max_raw - min_raw
                if raw_range < 1e-6:
                    raw_range = 1.0
                    
                norm_rerank_scores = [(s - min_raw) / raw_range for s in raw_scores]
                
                for idx, cid in enumerate(candidate_ids):
                    for item in candidates_to_rerank:
                        if item["id"] == cid:
                            # Gán lại score bằng reranker normalized score
                            item["score"] = norm_rerank_scores[idx]
                            break
        except Exception as e:
            print(f"⚠️ Lỗi trong quá trình chạy Reranker: {e}")

    # Gom nhóm tất cả kết quả
    all_candidates = candidates_to_rerank + remaining_candidates

    # Lấy metadata phục vụ Boosting
    doc_ids = list({item["doc_id"] for item in all_candidates})
    doc_meta = {}
    if doc_ids:
        placeholders = ",".join(["?"] * len(doc_ids))
        try:
            cursor.execute(f"SELECT id, so_ky_hieu, loai_van_ban, tinh_trang_hieu_luc, title FROM documents WHERE id IN ({placeholders})", doc_ids)
            for row in cursor.fetchall():
                doc_meta[row["id"]] = dict(row)
        except Exception as e:
            print(f"⚠️ Lỗi lấy metadata để boost: {e}")

    # Áp dụng Boosting vào score
    for item in all_candidates:
        doc_id = item["doc_id"]
        meta = doc_meta.get(doc_id)
        if not meta:
            continue
            
        so_ky_hieu = (meta.get("so_ky_hieu") or "").lower()
        loai_van_ban = (meta.get("loai_van_ban") or "").lower()
        status_str = (meta.get("tinh_trang_hieu_luc") or "").lower()

        # A. Flexible Symbol Match Boost
        symbol_boost = 0.0
        for sym in extracted_symbols:
            if sym in so_ky_hieu:
                symbol_boost = 3.0
                break
        if symbol_boost == 0.0 and numbers_in_q:
            nums_in_doc = re.findall(r'\b\d+\b', so_ky_hieu)
            for num_q in numbers_in_q:
                if num_q in nums_in_doc:
                    doc_type_short = loai_van_ban.replace("luật", "").replace("nghị định", "").strip()
                    if doc_type_short and doc_type_short in q_lower:
                        symbol_boost = 2.5
                        break
                    elif loai_van_ban in q_lower:
                        symbol_boost = 2.5
                        break
                    else:
                        symbol_boost = 1.0
                        break

        # B. Dynamic Document Type Boost
        type_boost = 0.0
        if is_detail_query:
            if "nghị định" in loai_van_ban:
                type_boost = 0.12
            elif "thông tư" in loai_van_ban:
                type_boost = 0.10
            elif "quyết định" in loai_van_ban:
                type_boost = 0.08
            elif "bộ luật" in loai_van_ban or "luật" in loai_van_ban:
                type_boost = 0.05
        elif is_general_query:
            if "hiến pháp" in loai_van_ban:
                type_boost = 0.20
            elif "bộ luật" in loai_van_ban or "luật" in loai_van_ban:
                type_boost = 0.15
            elif "pháp lệnh" in loai_van_ban:
                type_boost = 0.10
            elif "nghị định" in loai_van_ban:
                type_boost = 0.05
        else:
            if "hiến pháp" in loai_van_ban:
                type_boost = 0.12
            elif "bộ luật" in loai_van_ban or "luật" in loai_van_ban:
                type_boost = 0.10
            elif "pháp lệnh" in loai_van_ban:
                type_boost = 0.08
            elif "nghị định" in loai_van_ban:
                type_boost = 0.06
            elif "quyết định" in loai_van_ban:
                type_boost = 0.04
            elif "thông tư" in loai_van_ban:
                type_boost = 0.02

        # C. Status Boost
        status_boost = 0.0
        if "còn hiệu lực" in status_str or "hết hiệu lực một phần" in status_str:
            status_boost = 0.01

        # Thực thi công thức: cộng dồn metadata boost
        item["score"] = item["score"] + symbol_boost + (type_boost * 0.05) + status_boost

    # Sắp xếp ứng viên sau Boosting
    sorted_items = sorted(all_candidates, key=lambda x: x["score"], reverse=True)

    # E. Syllable Bigram Overlap Boost
    def get_unique_bigrams(text: str) -> set:
        if not text:
            return set()
        text = normalize_spelling(text)
        text = re.sub(r'[^\w\s]', ' ', text).strip()
        words = [w.lower() for w in text.split() if w]
        bigrams = set()
        for i in range(len(words) - 1):
            bigrams.add((words[i], words[i+1]))
        return bigrams

    def get_unique_words(text: str) -> set:
        if not text:
            return set()
        text = normalize_spelling(text)
        text = re.sub(r'[^\w\s]', ' ', text).strip()
        words = [w.lower() for w in text.split() if w]
        return set([w for w in words if w not in VIETNAMESE_STOPWORDS and len(w) > 1])

    q_words_clean = get_unique_words(q_processed)
    q_bigrams = get_unique_bigrams(q_processed)
    if q_words_clean and sorted_items:
        candidates_to_boost = sorted_items[:80]
        remaining_items = sorted_items[80:]
        
        candidate_ids = [item["id"] for item in candidates_to_boost]
        candidate_chunks = {}
        if candidate_ids:
            placeholders = ",".join(["?"] * len(candidate_ids))
            try:
                cursor.execute(f"SELECT id, chunk_text FROM document_chunks WHERE id IN ({placeholders})", candidate_ids)
                for row in cursor.fetchall():
                    candidate_chunks[row["id"]] = row["chunk_text"]
            except Exception as e:
                print(f"⚠️ Lỗi lấy chunk text để boost: {e}")
                
            for item in candidates_to_boost:
                cid = item["id"]
                chunk_text = candidate_chunks.get(cid, "")
                if chunk_text:
                    if len(q_bigrams) >= 2:
                        chunk_bigrams = get_unique_bigrams(chunk_text)
                        common_bigrams = q_bigrams & chunk_bigrams
                        overlap_ratio = len(common_bigrams) / len(q_bigrams) if q_bigrams else 0.0
                        
                        if overlap_ratio >= 0.70:
                            item["score"] += 3.5 * (overlap_ratio ** 2)
                        elif overlap_ratio >= 0.50:
                            item["score"] += 1.5 * overlap_ratio
                    else:
                        chunk_words = get_unique_words(chunk_text)
                        common = q_words_clean & chunk_words
                        overlap_ratio = len(common) / len(q_words_clean) if q_words_clean else 0.0
                        
                        if overlap_ratio >= 0.70:
                            item["score"] += 3.5 * (overlap_ratio ** 2)
                        elif overlap_ratio >= 0.50:
                            item["score"] += 1.5 * overlap_ratio
                            
        sorted_items = sorted(candidates_to_boost + remaining_items, key=lambda x: x["score"], reverse=True)

    # F. Hình phạt tài liệu địa phương & hết hiệu lực ở CUỐI CÙNG (Multiplicative Penalties)
    local_doc_penalty = 0.5
    inactive_doc_penalty = 0.8

    for item in sorted_items:
        doc_id = item["doc_id"]
        meta = doc_meta.get(doc_id)
        if not meta:
            continue
        
        so_ky_hieu_doc = (meta.get("so_ky_hieu") or "").lower()
        status_str_doc = (meta.get("tinh_trang_hieu_luc") or "").lower()
        co_quan_doc = (meta.get("co_quan_ban_hanh") or "").lower()
        title_doc = (meta.get("title") or "").lower()

        is_local_doc = False
        if "ubnd" in co_quan_doc or "hđnd" in co_quan_doc or "ủy ban nhân dân" in co_quan_doc or "hội đồng nhân dân" in co_quan_doc or "sở " in co_quan_doc or "sở" == co_quan_doc or "/qđ-ubnd" in so_ky_hieu_doc or "/nq-hđnd" in so_ky_hieu_doc:
            is_local_doc = True
        else:
            for p in PROVINCES:
                if p in co_quan_doc or p in title_doc:
                    is_local_doc = True
                    break
                    
        multiplier = 1.0
        if is_local_doc and not has_locality_in_q:
            multiplier *= local_doc_penalty
            
        if "hết hiệu lực toàn bộ" in status_str_doc:
            has_symbol_match = False
            for sym in extracted_symbols:
                if sym in so_ky_hieu_doc:
                    has_symbol_match = True
                    break
            if not has_symbol_match:
                multiplier *= inactive_doc_penalty
                
        item["score"] = item["score"] * multiplier

    # Sắp xếp lại sau hình phạt multiplicative
    sorted_items = sorted(sorted_items, key=lambda x: x["score"], reverse=True)

    # G. Diversification (Document Penalty)
    seen_docs_in_ranking = {}
    diversified_items = []
    for item in sorted_items:
        doc_id = item["doc_id"]
        if doc_id in seen_docs_in_ranking:
            seen_docs_in_ranking[doc_id] += 1
            penalty = 0.7 ** seen_docs_in_ranking[doc_id]
            item["score"] = item["score"] * penalty
        else:
            seen_docs_in_ranking[doc_id] = 0
        diversified_items.append(item)

    # Sắp xếp lại sau diversification
    sorted_items = sorted(diversified_items, key=lambda x: x["score"], reverse=True)
            
    total = len(sorted_items)

    # Phân trang
    paginated_items = sorted_items[offset:offset + limit]

    results = []
    if paginated_items:
        paginated_cids = [item["id"] for item in paginated_items]
        placeholders = ",".join(["?"] * len(paginated_cids))
        
        query_details = f"""
            SELECT c.id, c.doc_id, c.chunk_index, c.chunk_type, c.chunk_header, c.chunk_text, c.token_estimate,
                   d.title as document_title, d.so_ky_hieu as document_so_ky_hieu, d.loai_van_ban as document_loai_van_ban,
                   d.co_quan_ban_hanh as document_co_quan_ban_hanh, d.tinh_trang_hieu_luc as document_tinh_trang_hieu_luc,
                   d.ngay_ban_hanh as document_ngay_ban_hanh
            FROM document_chunks c
            JOIN documents d ON c.doc_id = d.id
            WHERE c.id IN ({placeholders})
        """
        try:
            cursor.execute(query_details, paginated_cids)
            details_map = {row["id"]: dict(row) for row in cursor.fetchall()}
            
            for item in paginated_items:
                cid = item["id"]
                if cid in details_map:
                    res = details_map[cid]
                    res["score"] = item["score"]
                    results.append(res)
        except Exception as e:
            print(f"⚠️ Lỗi lấy chi tiết chunks: {e}")

    conn.close()

    total_pages = (total + limit - 1) // limit if limit > 0 else 0
    current_page = (offset // limit) + 1 if limit > 0 else 1
    has_next = current_page < total_pages
    has_previous = current_page > 1

    return {
        "total": total,
        "limit": limit,
        "offset": offset,
        "total_pages": total_pages,
        "current_page": current_page,
        "has_next": has_next,
        "has_previous": has_previous,
        "results": results,
    }



@router.post("/reload-index", tags=["🔍 Tìm kiếm & Tra cứu (Luật)"], summary="Hot-reload BM25 index")
async def reload_index(_api_key: str = Depends(require_api_key)):
    """
    Hot-reload BM25 index từ cache mới nhất (sau khi auto_rebuild_index.py chạy).
    Không cần restart server.
    """
    import threading
    from app.hybrid_search import init_hybrid_engine
    from app.config import CONTENT_DB, DB_NAME
    
    def _reload():
        try:
            init_hybrid_engine(DB_NAME, CONTENT_DB)
        except Exception as e:
            print(f"⚠️ Reload failed: {e}")
    
    threading.Thread(target=_reload, daemon=True).start()
    
    return {"status": "reloading", "message": "BM25 index đang được reload trong background. Sẽ sẵn sàng trong ~10s."}


# ─────────────────── CATEGORIES ───────────────────

@simple_ttl_cache(ttl_seconds=3600)
def _get_cached_document_types():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT loai_van_ban, count(*) as cnt FROM documents GROUP BY loai_van_ban ORDER BY cnt DESC")
    rows = cursor.fetchall()
    conn.close()
    return [{"name": row[0] or "Khác", "count": row[1]} for row in rows]


@router.get("/categories/types", response_model=List[CategoryItem], tags=["🏷️ Danh mục (Luật)"], summary="Loại văn bản")
def get_document_types(_key=Depends(require_api_key)):
    """Danh sách loại văn bản. **Yêu cầu API Key.**"""
    return _get_cached_document_types()


@simple_ttl_cache(ttl_seconds=3600)
def _get_cached_fields():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT linh_vuc, count(*) as cnt FROM documents WHERE linh_vuc IS NOT NULL AND linh_vuc != '' GROUP BY linh_vuc ORDER BY cnt DESC"
    )
    rows = cursor.fetchall()
    conn.close()
    return [{"name": row[0], "count": row[1]} for row in rows]


@router.get("/categories/fields", response_model=List[CategoryItem], tags=["🏷️ Danh mục (Luật)"], summary="Lĩnh vực")
def get_fields(_key=Depends(require_api_key)):
    """Danh sách lĩnh vực pháp luật. **Yêu cầu API Key.**"""
    return _get_cached_fields()


@simple_ttl_cache(ttl_seconds=3600)
def _get_cached_agencies():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT co_quan_ban_hanh, count(*) as cnt FROM documents WHERE co_quan_ban_hanh IS NOT NULL AND co_quan_ban_hanh != '' GROUP BY co_quan_ban_hanh ORDER BY cnt DESC"
    )
    rows = cursor.fetchall()
    conn.close()
    return [{"name": row[0], "count": row[1]} for row in rows]


@router.get("/categories/agencies", response_model=List[CategoryItem], tags=["🏷️ Danh mục (Luật)"], summary="Cơ quan ban hành")
def get_agencies(_key=Depends(require_api_key)):
    """Danh sách cơ quan ban hành. **Yêu cầu API Key.**"""
    return _get_cached_agencies()


# ─────────────────── ADMINISTRATIVE DIVISIONS ───────────────────

@router.get("/provinces", response_model=List[ProvinceItem], tags=["🏷️ Danh mục (Luật)"], summary="Danh sách tỉnh/thành phố")
def get_provinces(_key=Depends(require_api_key)):
    """Lấy danh sách các tỉnh thành phố trực thuộc trung ương."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT code, name, full_name, code_name, administrative_unit_id FROM provinces ORDER BY name ASC")
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


@router.get("/provinces/{province_code}/wards", response_model=List[WardItem], tags=["🏷️ Danh mục (Luật)"], summary="Danh sách quận huyện/phường xã theo tỉnh")
def get_wards(
    province_code: str = Path(..., description="Mã tỉnh thành"),
    _key=Depends(require_api_key),
):
    """Lấy danh sách các đơn vị cấp dưới (quận/huyện/thị xã/phường/xã) của một tỉnh thành."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT code, name, full_name, code_name, province_code, administrative_unit_id "
        "FROM wards WHERE province_code = ? ORDER BY name ASC",
        (province_code,),
    )
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


# ─────────────────── DETAIL ───────────────────

@router.get("/{law_id}", response_model=LawDetail, summary="Chi tiết văn bản")
def get_law_detail(
    law_id: int = Path(..., description="ID văn bản"),
    _key=Depends(require_api_key),
):
    """Lấy toàn văn HTML và metadata đầy đủ. **Yêu cầu API Key.**"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM documents WHERE id = ?", (law_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail=f"Không tìm thấy văn bản có ID {law_id}")

    result = dict(row)

    # Nếu content_html đã được tách sang content_store.db
    if not result.get("content_html"):
        content_conn = get_content_connection()
        if content_conn:
            content_cursor = content_conn.cursor()
            content_cursor.execute(
                "SELECT content_html FROM document_content WHERE doc_id = ?",
                (law_id,),
            )
            content_row = content_cursor.fetchone()
            content_conn.close()
            if content_row:
                result["content_html"] = content_row["content_html"]

    return result


# ─────────────────── RELATIONSHIPS ───────────────────

def get_doc_level(loai_van_ban: Optional[str]) -> int:
    if not loai_van_ban:
        return 99
    l = loai_van_ban.lower().strip()
    if "hiến pháp" in l:
        return 1
    if "bộ luật" in l or "luật" in l:
        return 2
    if "pháp lệnh" in l:
        return 3
    if "lệnh" in l:
        return 4
    if "nghị định" in l:
        return 5
    if "thông tư" in l:
        return 6
    if "quyết định" in l:
        return 7
    if "chỉ thị" in l:
        return 8
    if "công văn" in l:
        return 9
    return 10


def is_local_doc(co_quan_ban_hanh: Optional[str], title: Optional[str], so_ky_hieu: Optional[str]) -> bool:
    for field in [co_quan_ban_hanh, title, so_ky_hieu]:
        if not field:
            continue
        f = field.lower()
        if "ủy ban nhân dân" in f or "ubnd" in f or "hội đồng nhân dân" in f or "hđnd" in f:
            return True
    return False


def classify_relationship(doc_a: dict, doc_b: dict) -> str:
    """
    doc_a: The cited document (our target, law_id)
    doc_b: The citing document (discovered from FTS match)
    """
    b_title = doc_b.get("title") or ""
    b_type = doc_b.get("loai_van_ban") or ""
    b_authority = doc_b.get("co_quan_ban_hanh") or ""
    b_symbol = doc_b.get("so_ky_hieu") or ""
    
    a_type = doc_a.get("loai_van_ban") or ""
    
    b_title_lower = b_title.lower()
    b_type_lower = b_type.lower().strip()
    a_type_lower = a_type.lower().strip()
    b_authority_lower = b_authority.lower()
    
    # 1. Local document check
    if is_local_doc(b_authority, b_title, b_symbol):
        rel = "Văn bản áp dụng địa phương (Quét Dọc)"
    # 2. Amendment check
    elif any(k in b_title_lower for k in ["sửa đổi", "bổ sung", "thay thế", "bãi bỏ"]):
        rel = "Văn bản sửa đổi, bổ sung, thay thế (Quét Ngang)"
    # 3. Cong van check
    elif "công văn" in b_type_lower:
        rel = "Công văn hướng dẫn nghiệp vụ (Quét Ngang)"
    # 4. Van ban hop nhat check
    elif "hợp nhất" in b_type_lower:
        rel = "Văn bản hợp nhất (Quét Ngang)"
    # 5. Supreme Court Judges' Council Resolution check
    elif "nghị quyết" in b_type_lower and ("hội đồng thẩm phán" in b_authority_lower or "tòa án nhân dân tối cao" in b_authority_lower):
        rel = "Nghị quyết HĐTP hướng dẫn áp dụng (Quét Ngang)"
    # 6. Systemic linkage check: both are major laws (Constitution, Code, Law)
    elif ("hiến pháp" in a_type_lower or "bộ luật" in a_type_lower or a_type_lower == "luật" or a_type_lower.startswith("luật ")) and \
         ("hiến pháp" in b_type_lower or "bộ luật" in b_type_lower or b_type_lower == "luật" or b_type_lower.startswith("luật ")):
        rel = "Liên kết đồng bộ hệ thống (Quét Ngang)"
    # 7. Specific vertical scans
    elif "nghị định" in b_type_lower and ("luật" in a_type_lower or "bộ luật" in a_type_lower or "hiến pháp" in a_type_lower):
        rel = "Nghị định hướng dẫn thi hành (Quét Dọc)"
    elif "thông tư" in b_type_lower and ("luật" in a_type_lower or "bộ luật" in a_type_lower or "hiến pháp" in a_type_lower):
        rel = "Thông tư hướng dẫn chi tiết (Quét Dọc)"
    elif "thông tư" in b_type_lower and "nghị định" in a_type_lower:
        rel = "Thông tư hướng dẫn thực hiện Nghị định (Quét Dọc)"
    elif "nghị quyết" in b_type_lower and ("luật" in a_type_lower or "bộ luật" in a_type_lower or "hiến pháp" in a_type_lower):
        rel = "Nghị quyết hướng dẫn thực hiện (Quét Dọc)"
    # 8. Fallback levels
    else:
        level_a = get_doc_level(a_type)
        level_b = get_doc_level(b_type)
        if level_b > level_a:
            rel = "Văn bản hướng dẫn, quy định chi tiết (Quét Dọc)"
        else:
            rel = "Văn bản dẫn chiếu, liên quan (Quét Ngang)"

    # Append status warning suffix
    b_status = (doc_b.get("tinh_trang_hieu_luc") or "").lower().strip()
    if "hết hiệu lực một phần" in b_status:
        rel += " (Hết hiệu lực một phần)"
    elif "hết hiệu lực" in b_status or "hết hiệu lực toàn bộ" in b_status:
        rel += " (Đã hết hiệu lực)"
        
    return rel


@router.get("/{law_id}/relationships", response_model=List[RelationshipInfo], tags=["🔗 Quan hệ pháp lý (Luật)"], summary="Quan hệ pháp lý")
def get_law_relationships(
    law_id: int = Path(..., description="ID văn bản"),
    _key=Depends(require_api_key),
):
    """Tra cứu mạng lưới liên kết pháp lý theo phương pháp Quét Dọc - Quét Ngang. **Yêu cầu API Key.**"""
    conn = get_db_connection()
    cursor = conn.cursor()

    # 1. Lấy metadata của văn bản hiện tại
    cursor.execute("SELECT id, title, so_ky_hieu, loai_van_ban, co_quan_ban_hanh, tinh_trang_hieu_luc FROM documents WHERE id = ?", (law_id,))
    doc_row = cursor.fetchone()
    if not doc_row:
        conn.close()
        raise HTTPException(status_code=404, detail=f"Không tìm thấy văn bản có ID {law_id}")
        
    doc_a = dict(doc_row)
    so_ky_hieu_a = doc_a.get("so_ky_hieu")

    # 2. Truy vấn các mối quan hệ cứng có sẵn trong DB
    query = """
        SELECT r.doc_id, r.other_doc_id, r.relationship,
               d.title as other_doc_title, d.so_ky_hieu as other_doc_so_ky_hieu,
               d.tinh_trang_hieu_luc as other_doc_tinh_trang_hieu_luc
        FROM relationships r
        JOIN documents d ON r.other_doc_id = d.id
        WHERE r.doc_id = ?
        UNION
        SELECT r.doc_id, r.other_doc_id, r.relationship,
               d.title as other_doc_title, d.so_ky_hieu as other_doc_so_ky_hieu,
               d.tinh_trang_hieu_luc as other_doc_tinh_trang_hieu_luc
        FROM relationships r
        JOIN documents d ON r.doc_id = d.id
        WHERE r.other_doc_id = ?
    """

    cursor.execute(query, (law_id, law_id))
    rows = cursor.fetchall()
    
    # Chuyển đổi thành list dict và ghi nhận related_ids đã có để deduplicate
    relationships_list = []
    seen_related_ids = set()
    
    for row in rows:
        r_dict = dict(row)
        
        # Append status warning suffix to hardcoded relationship label if the other doc is expired
        other_status = (r_dict.get("other_doc_tinh_trang_hieu_luc") or "").lower().strip()
        status_suffix = ""
        if "hết hiệu lực một phần" in other_status:
            status_suffix = " (Hết hiệu lực một phần)"
        elif "hết hiệu lực" in other_status or "hết hiệu lực toàn bộ" in other_status:
            status_suffix = " (Đã hết hiệu lực)"
            
        if status_suffix:
            r_dict["relationship"] = r_dict["relationship"] + status_suffix
            
        relationships_list.append(r_dict)
        
        # Ghi nhận ID liên kết để tránh trùng lặp
        if r_dict["doc_id"] == law_id:
            seen_related_ids.add(r_dict["other_doc_id"])
        else:
            seen_related_ids.add(r_dict["doc_id"])

    # 3. Quét FTS5 tìm văn bản dẫn chiếu ẩn trong nội dung toàn văn
    # Bỏ qua nếu là văn bản "Không số" hoặc trống để tránh false positive hàng loạt
    if so_ky_hieu_a and "không số" not in so_ky_hieu_a.lower() and so_ky_hieu_a.strip():
        tokens = re.findall(r'\w+', so_ky_hieu_a)
        if tokens:
            phrase = " ".join(tokens)
            fts_query = f'"{phrase}"'
            
            # Query content_fts tìm các văn bản dẫn chiếu (giới hạn 2000 dòng cho an toàn)
            cursor.execute("""
                SELECT d.id, d.title, d.so_ky_hieu, d.loai_van_ban, d.co_quan_ban_hanh, d.tinh_trang_hieu_luc
                FROM content_fts f
                JOIN documents d ON f.rowid = d.id
                WHERE f.content_text MATCH ?
                LIMIT 2000
            """, (fts_query,))
            citing_docs = cursor.fetchall()
            
            for doc_row_b in citing_docs:
                doc_b = dict(doc_row_b)
                b_id = doc_b["id"]
                
                # Bỏ qua nếu trùng chính nó hoặc đã có trong quan hệ cứng
                if b_id == law_id or b_id in seen_related_ids:
                    continue
                    
                seen_related_ids.add(b_id)
                
                # Phân loại mối quan hệ theo phương pháp Quét Dọc - Quét Ngang
                rel_label = classify_relationship(doc_a, doc_b)
                
                # Thêm vào kết quả
                relationships_list.append({
                    "doc_id": b_id,
                    "other_doc_id": law_id,
                    "relationship": rel_label,
                    "other_doc_title": doc_b["title"],
                    "other_doc_so_ky_hieu": doc_b["so_ky_hieu"]
                })

    conn.close()
    return relationships_list


@router.get("/{law_id}/article-modifications", response_model=List[ArticleModification], tags=["🔗 Quan hệ pháp lý (Luật)"], summary="Các điều khoản bị sửa đổi")
def get_article_modifications(
    law_id: int = Path(..., description="ID văn bản"),
    _key=Depends(require_api_key),
):
    """Tra cứu các điều khoản bị sửa đổi của văn bản. **Yêu cầu API Key.**"""
    conn = get_db_connection()
    cursor = conn.cursor()

    query = """
        SELECT article_name, modified_text, modified_by_doc_id
        FROM article_modifications
        WHERE doc_id = ?
    """

    try:
        cursor.execute(query, (law_id,))
        rows = cursor.fetchall()
    except sqlite3.OperationalError:
        # Table might not exist yet if script hasn't run fully
        rows = []

    conn.close()
    return [dict(row) for row in rows]


@router.get("/{law_id}/download", summary="Tải file DOCX của văn bản")
def download_law_docx(
    law_id: int = Path(..., description="ID văn bản"),
    _key=Depends(require_api_key),
):
    """
    Tải văn bản pháp luật dưới dạng file Microsoft Word (.docx).
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM documents WHERE id = ?", (law_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail=f"Không tìm thấy văn bản có ID {law_id}")

    metadata = dict(row)
    title = metadata.get("title") or "Van ban phap luat"

    # Lấy content_html
    content_html = ""
    content_conn = get_content_connection()
    if content_conn:
        content_cursor = content_conn.cursor()
        content_cursor.execute(
            "SELECT content_html FROM document_content WHERE doc_id = ?",
            (law_id,),
        )
        content_row = content_cursor.fetchone()
        content_conn.close()
        if content_row:
            content_html = content_row["content_html"]

    # Parse and build docx
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 1. Quốc hiệu / Tiêu ngữ
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_header1 = p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_header1.bold = True
    r_header1.font.size = Pt(12)
    r_header2 = p_header.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r_header2.bold = True
    r_header2.font.size = Pt(11)
    r_header3 = p_header.add_run("---------------")
    r_header3.font.size = Pt(11)

    # 2. Metadata (Cơ quan & Số hiệu)
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    organ = metadata.get("co_quan_ban_hanh") or "CƠ QUAN BAN HÀNH"
    so_ky_hieu = metadata.get("so_ky_hieu") or "N/A"
    ngay_ban_hanh = metadata.get("ngay_ban_hanh") or "N/A"
    
    p_meta.add_run(f"{organ.upper()}\nSố: {so_ky_hieu}\n").bold = True
    
    # 3. Ngày ban hành (Căn phải)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"Hà Nội, ngày {ngay_ban_hanh}").italic = True

    doc.add_paragraph()  # Spacing

    # 4. Tiêu đề chính
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title.upper())
    r_title.bold = True
    r_title.font.size = Pt(13)

    doc.add_paragraph()  # Spacing

    # 5. Nội dung văn bản
    if content_html:
        soup = BeautifulSoup(content_html, "html.parser")
        content_container = soup.find(id="content") or soup.find(class_="noi-dung") or soup
        
        # Traverse elements
        for block in content_container.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "table", "li"]):
            if block.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(block.get_text().strip())
                r.bold = True
                if block.name == "h1":
                    r.font.size = Pt(13)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    r.font.size = Pt(12)
            elif block.name == "p":
                if not block.find("table"):
                    text = block.get_text().strip()
                    if text:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.15
                        p.add_run(text)
            elif block.name == "li":
                text = block.get_text().strip()
                if text:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(3)
                    p.add_run(text)
            elif block.name == "table":
                rows = block.find_all("tr")
                if rows:
                    max_cols = max(len(r.find_all(["td", "th"])) for r in rows)
                    if max_cols > 0:
                        tbl = doc.add_table(rows=0, cols=max_cols)
                        tbl.style = 'Table Grid'
                        for r_el in rows:
                            cells = r_el.find_all(["td", "th"])
                            row_cells = tbl.add_row().cells
                            for c_idx, cell in enumerate(cells):
                                if c_idx < len(row_cells):
                                    row_cells[c_idx].text = cell.get_text().strip()
    else:
        p = doc.add_paragraph()
        p.add_run("Văn bản chưa cập nhật nội dung toàn văn.")

    # Save to Stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Safe file name
    safe_title = "".join([c if c.isalnum() or c in "._-" else "_" for c in so_ky_hieu])
    if not safe_title:
        safe_title = f"document_{law_id}"
    
    import urllib.parse
    encoded_filename = urllib.parse.quote(f"{safe_title}.docx")
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )







