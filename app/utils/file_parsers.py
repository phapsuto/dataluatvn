import fitz  # PyMuPDF
from docx import Document
import io
import requests
import base64
import os
import subprocess
import tempfile

def _tesseract_ocr_image_bytes(img_bytes: bytes) -> str:
    """Fallback local OCR using Tesseract CLI with Vietnamese (vie) language."""
    try:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(img_bytes)
            tmp_path = tmp.name
        
        proc = subprocess.run(
            ["tesseract", tmp_path, "stdout", "-l", "vie"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=30
        )
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
            
        if proc.returncode == 0 and proc.stdout:
            return proc.stdout.strip()
    except Exception as e:
        print(f"[Tesseract Fallback Error] {e}")
    return ""

def _fpt_ocr_image_base64(base64_img: str, img_bytes: bytes = None) -> str:
    """Calls FPT Cloud Vision API (Qwen2.5-VL-7B-Instruct) to extract text from a base64 image.
    Falls back to local Tesseract OCR if API fails or key is missing.
    """
    api_key = os.environ.get("FPT_CLOUD_API_KEY") or os.environ.get("FPT_API_KEY")
    if api_key:
        url = "https://mkp-api.fptcloud.com/v1/chat/completions"
        payload = {
            "model": "Qwen2.5-VL-7B-Instruct",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Trích xuất toàn bộ văn bản và bảng biểu có trong hình ảnh này. Chỉ trả về văn bản được trích xuất (nếu có bảng biểu hãy trình bày dạng Markdown), không giải thích gì thêm."},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_img}"}}
                    ]
                }
            ],
            "temperature": 0.1,
            "max_tokens": 2048
        }
        try:
            res = requests.post(url, json=payload, headers={"Authorization": f"Bearer {api_key}"}, timeout=60)
            if res.status_code == 200:
                data = res.json()
                content = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
                if content:
                    return content
        except Exception as e:
            print(f"[FPT Vision OCR Error, falling back to Tesseract] {e}")

    # Fallback to Tesseract OCR if FPT API fails or key missing
    if img_bytes:
        return _tesseract_ocr_image_bytes(img_bytes)
    else:
        try:
            return _tesseract_ocr_image_bytes(base64.b64decode(base64_img))
        except Exception:
            return ""

def parse_image(file_bytes: bytes, filename: str, progress_callback=None) -> str:
    """Extracts text and layout from an image file (.png, .jpg, .webp, .bmp)."""
    if progress_callback:
        progress_callback(1, 1)
    b64_img = base64.b64encode(file_bytes).decode("utf-8")
    text = _fpt_ocr_image_base64(b64_img, file_bytes)
    if not text.strip():
        raise ValueError(f"Không thể nhận diện chữ trong ảnh {filename} (OCR trống).")
    return text.strip()

def parse_pdf(file_bytes: bytes, progress_callback=None) -> str:
    """Extracts text from a PDF file. Uses FPT Vision OCR for scanned pages."""
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        page_results = [None] * total_pages
        scanned_pages = []
        completed = 0
        
        for i, page in enumerate(doc):
            page_text = page.get_text().strip()
            if len(page_text) < 50:
                scanned_pages.append(i)
            else:
                page_results[i] = page_text + "\n\n"
                completed += 1
                if progress_callback:
                    progress_callback(completed, total_pages)
        
        # Run OCR in batches to avoid OOM for large documents
        BATCH_SIZE = 5
        if scanned_pages:
            for batch_start in range(0, len(scanned_pages), BATCH_SIZE):
                batch_indices = scanned_pages[batch_start:batch_start + BATCH_SIZE]
                tasks = []
                
                # Generate images only for the current batch
                for i in batch_indices:
                    pix = doc[i].get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")
                    b64_img = base64.b64encode(img_bytes).decode("utf-8")
                    tasks.append((i, b64_img, img_bytes))
                
                # Process the batch concurrently
                with ThreadPoolExecutor(max_workers=BATCH_SIZE) as executor:
                    future_to_idx = {executor.submit(_fpt_ocr_image_base64, b64, ibytes): idx for idx, b64, ibytes in tasks}
                    for future in as_completed(future_to_idx):
                        idx = future_to_idx[future]
                        try:
                            ocr_text = future.result()
                            if not ocr_text.strip():
                                page_results[idx] = "\n\n[LỖI: TRANG NÀY KHÔNG NHẬN DIỆN ĐƯỢC CHỮ]\n\n"
                            else:
                                page_results[idx] = ocr_text + "\n\n"
                        except Exception as e:
                            print(f"[OCR Task Error] {e}")
                            page_results[idx] = "\n\n[LỖI: QUÁ TRÌNH OCR TRANG NÀY BỊ THẤT BẠI]\n\n"
                        
                        completed += 1
                        if progress_callback:
                            progress_callback(completed, total_pages)
                        
        text = "".join(filter(None, page_results))
    except Exception as e:
        raise ValueError(f"Lỗi khi đọc file PDF: {str(e)}")
    return text.strip()

def _table_to_markdown(table) -> str:
    """Converts a python-docx table to a Markdown formatted table."""
    lines = []
    headers = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if i == 0:
            headers = cells
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        else:
            lines.append("| " + " | ".join(cells) + " |")
    return "\n" + "\n".join(lines) + "\n\n"

def parse_docx(file_bytes: bytes, progress_callback=None) -> str:
    """Extracts text and tables (formatted as Markdown) from a DOCX file."""
    text_parts = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        # Extract paragraphs
        for para in doc.paragraphs:
            val = para.text.strip()
            if val:
                text_parts.append(val)
        # Extract tables as Markdown
        if doc.tables:
            text_parts.append("\n--- [BẢNG BIỂU CHI TIẾT TRONG TÀI LIỆU] ---")
            for table in doc.tables:
                text_parts.append(_table_to_markdown(table))
        if progress_callback:
            progress_callback(1, 1)
    except Exception as e:
        raise ValueError(f"Lỗi khi đọc file DOCX: {str(e)}")
    return "\n\n".join(text_parts).strip()

def parse_txt(file_bytes: bytes, progress_callback=None) -> str:
    """Extracts text from a TXT file."""
    try:
        if progress_callback:
            progress_callback(1, 1)
        return file_bytes.decode('utf-8').strip()
    except UnicodeDecodeError:
        try:
            return file_bytes.decode('windows-1258').strip()
        except Exception:
            raise ValueError("Lỗi khi đọc file TXT: Không thể giải mã định dạng file")
    except Exception as e:
        raise ValueError(f"Lỗi khi đọc file TXT: {str(e)}")

def parse_file(filename: str, file_bytes: bytes, progress_callback=None) -> str:
    """Detects file type and extracts text + tables."""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return parse_pdf(file_bytes, progress_callback)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        if progress_callback:
            progress_callback(1, 1)
        return parse_docx(file_bytes, progress_callback)
    elif filename_lower.endswith('.txt') or filename_lower.endswith('.csv'):
        return parse_txt(file_bytes, progress_callback)
    elif any(filename_lower.endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff']):
        return parse_image(file_bytes, filename, progress_callback)
    else:
        raise ValueError("Định dạng file không được hỗ trợ. Chỉ hỗ trợ PDF, DOCX, DOC, TXT, CSV và Ảnh (PNG, JPG, WEBP).")

